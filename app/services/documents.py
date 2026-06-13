"""Extract plain text from uploaded travel documents (PDF, DOCX, TXT, MD)."""
from __future__ import annotations

import io
from pathlib import PurePosixPath


class DocumentError(ValueError):
    pass


def extract_text(filename: str, data: bytes) -> str:
    suffix = PurePosixPath(filename.lower()).suffix
    if suffix == ".pdf":
        text = _from_pdf(data)
    elif suffix == ".docx":
        text = _from_docx(data)
    elif suffix in {".txt", ".md", ".markdown", ".text"}:
        text = data.decode("utf-8", errors="replace")
    else:
        raise DocumentError(f"Unsupported file type '{suffix or filename}'. Use PDF, DOCX, TXT or Markdown.")

    text = text.strip()
    if len(text) < 20:
        raise DocumentError("Could not extract readable text from the document (is it a scanned image?).")
    return text


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise DocumentError(f"Could not parse PDF: {exc}") from exc
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _from_docx(data: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise DocumentError(f"Could not parse DOCX: {exc}") from exc
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)
